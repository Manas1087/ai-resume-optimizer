import streamlit as st
import fitz  # PyMuPDF
from google import genai
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import re
import os

# ===============================
# 🎨 UI STYLE
# ===============================
st.markdown("""
<style>
body {
    background-color: #0f172a;
}
h1, h2, h3 {
    color: #e2e8f0;
}
.stButton>button {
    background-color: #2563eb;
    color: white;
    border-radius: 10px;
    padding: 10px 20px;
    font-weight: bold;
}
</style>
""", unsafe_allow_html=True)

# ===============================
# 🔑 GEMINI CLIENT (SECURE)
# ===============================
client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
MODEL = "gemini-2.5-flash"

# ===============================
# 📄 Extract PDF
# ===============================
def extract_text_from_pdf(uploaded_file):
    text = ""
    with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

# ===============================
# 🤖 ANALYZE
# ===============================
def analyze_resume(jd, resume):
    prompt = f"""
Analyze the resume against the job description.

STRICT FORMAT:

Match Score: <number>%

Matching Skills:
- skill1
- skill2

Missing Skills:
- skill1
- skill2

Suggestions:
- point1
- point2

Job Description:
{jd}

Resume:
{resume}
"""
    response = client.models.generate_content(
        model=MODEL,
        contents=prompt
    )
    return response.text

# ===============================
# ✨ OPTIMIZE
# ===============================
def optimize_resume(jd, resume):
    prompt = f"""
Rewrite the resume keeping structure same.

- Use bullet points
- ATS friendly
- Add missing keywords

Return only formatted resume.

JD:
{jd}

Resume:
{resume}
"""
    response = client.models.generate_content(
        model=MODEL,
        contents=prompt
    )
    return response.text

# ===============================
# 📄 DOCX FORMAT
# ===============================
def create_docx(content):
    doc = Document()

    for i, line in enumerate(content.split("\n")):
        line = line.strip()
        if not line:
            continue

        if i == 0:
            run = doc.add_paragraph().add_run(line)
            run.bold = True
            run.font.size = Pt(16)

        elif line.upper() in ["SUMMARY", "SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION"]:
            run = doc.add_paragraph().add_run(line.upper())
            run.bold = True

        elif line.startswith(("•", "-", "*")):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')

        else:
            doc.add_paragraph(line)

    doc.save("optimized_resume.docx")

# ===============================
# 🔁 PDF
# ===============================
def convert_to_pdf():
    try:
        convert("optimized_resume.docx", "optimized_resume.pdf")
        return True
    except:
        return False

# ===============================
# 🧠 PARSE ANALYSIS
# ===============================
def parse_analysis(text):
    sections = {
        "score": "",
        "matching": [],
        "missing": [],
        "suggestions": []
    }

    current = None

    for line in text.split("\n"):
        line = line.strip()

        if "Match Score" in line:
            sections["score"] = line.split(":")[1].strip()

        elif line.startswith("Matching Skills"):
            current = "matching"

        elif line.startswith("Missing Skills"):
            current = "missing"

        elif line.startswith("Suggestions"):
            current = "suggestions"

        elif line.startswith("-") and current:
            sections[current].append(line[1:].strip())

    return sections

# ===============================
# 🌐 UI
# ===============================
st.title("🚀 AI Resume Optimizer")

uploaded_file = st.file_uploader("Upload Resume (PDF)", type=["pdf"])
jd = st.text_area("Paste Job Description")

# ===============================
# 🔍 ANALYZE
# ===============================
if uploaded_file and jd:
    if st.button("🔍 Analyze Resume"):

        resume_text = extract_text_from_pdf(uploaded_file)

        with st.spinner("Analyzing..."):
            analysis = analyze_resume(jd, resume_text)

        st.session_state["analysis"] = analysis
        st.session_state["resume_text"] = resume_text

# ===============================
# 📊 DISPLAY ANALYSIS
# ===============================
if "analysis" in st.session_state:

    data = parse_analysis(st.session_state["analysis"])
    resume_text = st.session_state["resume_text"]

    st.subheader("📊 Resume Analysis")

    # Score
    score_text = data["score"]
    score = int(re.search(r"\d+", score_text).group()) if score_text else 0

    col1, col2, col3 = st.columns(3)
    col1.metric("🎯 Match Score", score_text)
    col2.metric("📈 Strength", "Strong" if score >= 70 else "Moderate")
    col3.metric("⚠️ Improve", "Yes" if score < 70 else "Low")

    st.progress(score / 100)

    # Matching Skills
    st.markdown("### ✅ Matching Skills")
    for skill in data["matching"]:
        st.markdown(f"- {skill}")

    # Missing Skills
    st.markdown("### ❌ Missing Skills")
    for skill in data["missing"]:
        st.markdown(f"- {skill}")

    # Suggestions
    st.markdown("### 💡 Suggestions")
    for s in data["suggestions"]:
        st.markdown(f"- {s}")

    # ===============================
    # 🚀 OPTIMIZE
    # ===============================
    if st.button("🚀 Generate Optimized Resume"):

        with st.spinner("Optimizing..."):
            optimized = optimize_resume(jd, resume_text)

        st.subheader("✨ Optimized Resume")

        st.text_area(
            "Preview",
            optimized,
            height=400,
            label_visibility="collapsed"
        )

        create_docx(optimized)

        if convert_to_pdf():
            with open("optimized_resume.pdf", "rb") as f:
                st.download_button("📥 Download PDF", f)
        else:
            with open("optimized_resume.docx", "rb") as f:
                st.download_button("📥 Download DOCX", f)
